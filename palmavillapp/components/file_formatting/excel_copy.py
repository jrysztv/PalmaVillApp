from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def get_rgb_color(color):
    if color.type == "rgb":
        return color.rgb[2:]  # Remove 'FF' prefix
    elif color.type == "theme":
        # Convert theme color to RGB (a simple approximation)
        # In practice, you may need to look up the theme color from the workbook
        theme_colors = {
            0: "FFFFFF",  # Default white for theme color 0
            # Add other theme color mappings as needed
        }
        return theme_colors.get(
            color.theme, "000000"
        )  # Default to black if theme color not found
    return None


# def add_thick_border(cell):
#     # Create a thick border around the cell
#     border = parse_xml(
#         r'<w:tcBorders %s><w:top w:val="single" w:sz="12"/><w:left w:val="single" w:sz="12"/><w:bottom w:val="single" w:sz="12"/><w:right w:val="single" w:sz="12"/></w:tcBorders>'
#         % nsdecls("w")
#     )
#     tc_pr = cell._element.get_or_add_tcPr()
#     tc_pr.append(border)


def add_border(cell, border_type="slim"):
    if border_type == "thick":
        border = parse_xml(
            r'<w:tcBorders %s><w:top w:val="single" w:sz="24"/><w:left w:val="single" w:sz="24"/><w:bottom w:val="single" w:sz="24"/><w:right w:val="single" w:sz="24"/></w:tcBorders>'
            % nsdecls("w")
        )
    else:  # 'slim'
        border = parse_xml(
            r'<w:tcBorders %s><w:top w:val="single" w:sz="6"/><w:left w:val="single" w:sz="6"/><w:bottom w:val="single" w:sz="6"/><w:right w:val="single" w:sz="6"/></w:tcBorders>'
            % nsdecls("w")
        )
    tc_pr = cell._element.get_or_add_tcPr()
    tc_pr.append(border)


def set_paragraph_format(paragraph):
    # Set the paragraph format to remove spacing and set line spacing to 1.0
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.line_spacing = Pt(12)  # Single line spacing


def copy_excel_to_word(workbook, word_path, table_index=0, doc=None):
    sheet = workbook.active

    if len(doc.tables) > table_index:
        table = doc.tables[table_index]
    else:
        raise ValueError(
            f"Table index {table_index} out of range. Document contains {len(doc.tables)} tables."
        )

    # Clear existing table contents
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""

    # Update the table to match Excel sheet dimensions
    if len(table.rows) < sheet.max_row:
        for _ in range(sheet.max_row - len(table.rows)):
            table.add_row()
    elif len(table.rows) > sheet.max_row:
        for _ in range(len(table.rows) - sheet.max_row):
            table.delete_row(table.rows[-1])

    if len(table.columns) < sheet.max_column:
        for _ in range(sheet.max_column - len(table.columns)):
            table.add_column()
    elif len(table.columns) > sheet.max_column:
        for _ in range(len(table.columns) - sheet.max_column):
            for row in table.rows:
                row.cells[-1]._element.getparent().remove(row.cells[-1]._element)

    # Set column widths based on Excel column widths
    for col_idx, col in enumerate(sheet.iter_cols(values_only=False)):
        column_width = sheet.column_dimensions[col[0].column_letter].width
        if column_width is not None:
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(column_width / 10)  # Adjust the factor as needed

    # Set row heights based on Excel row heights
    for row_idx, row in enumerate(sheet.iter_rows(values_only=False)):
        row_height = sheet.row_dimensions[row_idx + 1].height
        if row_height is not None:
            table.rows[row_idx].height = Inches(
                row_height / 70
            )  # Adjust the factor as needed

    # Copy Excel sheet data to the Word table
    for i, row in enumerate(sheet.iter_rows(values_only=False)):
        for j, cell in enumerate(row):
            word_cell = table.cell(i, j)
            word_cell.text = str(cell.value).strip() if cell.value is not None else ""

            # Add thick border to the cell
            add_border(word_cell)

            # Set paragraph formatting
            set_paragraph_format(word_cell.paragraphs[0])

            # Copy font formatting
            if cell.font:
                word_font = word_cell.paragraphs[0].runs[0].font
                if cell.font.bold:
                    word_font.bold = True
                if cell.font.italic:
                    word_font.italic = True
                if cell.font.size:
                    word_font.size = Pt(cell.font.size)
                if cell.font.name:
                    word_font.name = cell.font.name
                if cell.font.color:
                    rgb_str = get_rgb_color(cell.font.color)
                    if rgb_str:
                        word_font.color.rgb = RGBColor(
                            int(rgb_str[0:2], 16),
                            int(rgb_str[2:4], 16),
                            int(rgb_str[4:6], 16),
                        )

            # Copy alignment
            if cell.alignment:
                word_cell.vertical_alignment = (
                    WD_ALIGN_VERTICAL.CENTER
                    if cell.alignment.vertical == "center"
                    else WD_ALIGN_VERTICAL.TOP
                )
                if cell.alignment.horizontal == "center":
                    word_cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                elif cell.alignment.horizontal == "right":
                    word_cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT

            # Copy fill color
            if cell.fill and cell.fill.start_color.index != "00000000":
                tc = word_cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                fill_rgb_str = get_rgb_color(cell.fill.start_color)
                if fill_rgb_str:
                    shd.set(qn("w:fill"), fill_rgb_str)
                tcPr.append(shd)

    # Save the Word document
    # doc.save(word_path)
