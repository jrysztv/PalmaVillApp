from docx import Document
import openpyxl
import pandas as pd
from palmavillapp.components.constants import FELHO_COLUMNS, TABLE_TEMPLATE_PATH


from palmavillapp.components.constants import IFA_TEMPLATE_PATH
from palmavillapp.components.data_pipelines.guest_nights_calculations import (
    calculate_guest_nights_by_age_month,
)
from palmavillapp.components.file_formatting.excel_copy import copy_excel_to_word


def insert_values_table(kid_nights, adult_nights):
    total_nights = kid_nights + adult_nights
    # Load the Excel template

    template_path = TABLE_TEMPLATE_PATH
    workbook = openpyxl.load_workbook(template_path)

    # Select the worksheet
    worksheet = workbook["Sheet1"]  # Replace 'Sheet1' with your sheet name

    # Insert values into specific cells
    worksheet["B2"] = f"{total_nights} db"
    worksheet["B4"] = f"{kid_nights} db"
    worksheet["B6"] = f"{kid_nights} db"
    worksheet["B17"] = f"{adult_nights} db"
    worksheet["B18"] = f"{adult_nights*530} Ft"

    return workbook


def update_template_word_with_values(kid_nights, adult_nights, year, month):
    total_ifa = adult_nights * 530
    table_workbook = insert_values_table(kid_nights, adult_nights)

    doc = Document(IFA_TEMPLATE_PATH)
    copy_excel_to_word(table_workbook, IFA_TEMPLATE_PATH, 4, doc=doc)
    # Update header with other info
    doc.tables[0].cell(2, 1).text = f"{total_ifa} Ft"
    doc.tables[0].cell(4, 1).text = f"{year} év {month.lower()}"
    return doc


if __name__ == "__main__":
    df_felho = pd.read_excel("data/felho.xlsx")
    df_felho.columns = FELHO_COLUMNS
    guest_nights = calculate_guest_nights_by_age_month(df_felho=df_felho)
    date_year = guest_nights["date_year"]
    guest_nights = guest_nights["monthly"]
    months_detected = list(guest_nights.keys())
    # example usage of generating a word document from an excel template
    month = months_detected[0]
    total_nights = guest_nights[month]["adult"] + guest_nights[month]["kid"]
    kid_nights = guest_nights[month]["kid"]
    adult_nights = guest_nights[month]["adult"]
    total_ifa = guest_nights[month]["adult"] * 530
    update_template_word_with_values(kid_nights, adult_nights, date_year, month)
